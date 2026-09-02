import json
import os
from pathlib import Path
import subprocess
import sys
import tempfile
import unittest

import yaml


ROOT = Path(__file__).resolve().parents[1]
CURSOR = ROOT / ".cursor"
CLIENT_CURSOR = ROOT / "adapters" / "cursor" / "templates" / ".cursor"


class CursorCliConfigurationTests(unittest.TestCase):
    def test_ui_and_cli_permission_profiles_coexist(self):
        ui_config = json.loads((CURSOR / "permissions.json").read_text(encoding="utf-8"))
        cli_config = json.loads((CURSOR / "cli.json").read_text(encoding="utf-8"))

        self.assertIn("autoRun", ui_config)
        self.assertEqual(set(cli_config), {"permissions"})

        allow = set(cli_config["permissions"]["allow"])
        deny = set(cli_config["permissions"]["deny"])
        self.assertIn("Shell(git:status*)", allow)
        self.assertIn("Shell(git:add*)", allow)
        self.assertIn("Shell(git:commit*)", allow)
        self.assertIn("Shell(python:scripts/ai-team/diagnose.py*)", allow)
        self.assertIn("Shell(python:scripts/ai-team/sync_source_manifest.py*)", allow)
        self.assertIn("Shell(python:scripts/ai-team/check_git_policy.py*)", allow)
        self.assertIn("Shell(py:-3 scripts/ai-team/validate.py*)", allow)
        self.assertIn("Shell(py:-3 scripts/ai-team/sync_source_manifest.py*)", allow)
        self.assertIn("Shell(py:-3 scripts/ai-team/check_git_policy.py*)", allow)
        # Themed read-only additions: git reconnaissance verbs beyond the
        # original status/diff/log/show, and read-only PowerShell cmdlets
        # (never a write cmdlet like Remove-Item/Set-Content).
        self.assertIn("Shell(git:rev-parse*)", allow)
        self.assertIn("Shell(git:merge-base*)", allow)
        self.assertIn("Shell(git:branch*)", allow)
        self.assertIn("Shell(Get-ChildItem:*)", allow)
        self.assertIn("Shell(Select-String:*)", allow)
        self.assertIn("Shell(Test-Path:*)", allow)
        for write_cmdlet in ("Remove-Item", "Set-Content", "Stop-Process", "New-Item"):
            self.assertFalse(
                any(entry.startswith(f"Shell({write_cmdlet}") for entry in allow),
                f"{write_cmdlet} must stay behind approval, not be broadly allowed",
            )
        self.assertIn("Write(.ai-team/constitution/**)", deny)
        self.assertIn("Write(.cursor/cli.json)", deny)
        self.assertIn("Write(.cursor/permissions.json)", deny)
        self.assertIn("Shell(git:reset*--hard*)", deny)
        self.assertIn("Shell(git:commit*--amend*)", deny)
        self.assertIn("Shell(git:rebase*)", deny)
        # The broadened Shell(git:branch*) allow must not open a path to
        # force-deleting a branch without approval.
        self.assertIn("Shell(git:branch*-D*)", deny)
        self.assertIn("Shell(git:branch*--delete*--force*)", deny)

        terminal_allowlist = set(ui_config.get("terminalAllowlist") or [])
        self.assertIn("git add", terminal_allowlist)
        self.assertIn("git commit", terminal_allowlist)
        self.assertIn("git rev-parse", terminal_allowlist)
        self.assertIn("git merge-base", terminal_allowlist)
        self.assertIn("Get-ChildItem", terminal_allowlist)
        self.assertIn("python scripts/ai-team/sync_source_manifest.py", terminal_allowlist)
        self.assertIn("python scripts/ai-team/check_git_policy.py", terminal_allowlist)

    def test_subagents_are_foreground_by_default_and_readonly_roles_stay_readonly(self):
        readonly_roles = {
            "architect.md",
            "auditor.md",
            "code-reviewer.md",
            "product-analyst.md",
            "release-agent.md",
            "security-reviewer.md",
        }
        for agent_path in sorted((CLIENT_CURSOR / "agents").glob("*.md")):
            text = agent_path.read_text(encoding="utf-8")
            self.assertNotIn("is_background: true", text, agent_path.name)
            if agent_path.name in readonly_roles:
                self.assertIn("readonly: true", text, agent_path.name)

    def test_auth_smoke_agent_is_non_readonly_for_cross_platform_allowlist_smoke(self):
        text = (CLIENT_CURSOR / "agents" / "auth-smoke.md").read_text(encoding="utf-8")
        self.assertIn("name: auth-smoke", text)
        self.assertIn("readonly: false", text)
        self.assertNotIn("readonly: true", text)
        self.assertIn("Cross-platform CLI Allowlist", text)

    def test_hooks_cover_shell_and_subagent_lifecycle(self):
        hooks = json.loads((CURSOR / "hooks.json").read_text(encoding="utf-8"))["hooks"]
        for hook_name in [
            "beforeShellExecution",
            "afterShellExecution",
            "subagentStart",
            "subagentStop",
        ]:
            self.assertIn(hook_name, hooks)
            self.assertTrue(hooks[hook_name])

    def test_hooks_use_portable_python_runner(self):
        hooks = json.loads((CURSOR / "hooks.json").read_text(encoding="utf-8"))["hooks"]
        commands = [item["command"] for definitions in hooks.values() for item in definitions]
        self.assertTrue(commands)
        for command in commands:
            self.assertTrue(
                command.startswith(".cursor/hooks/run_hook.cmd "), command
            )

        runner = (CURSOR / "hooks" / "run_hook.cmd").read_text(encoding="utf-8")
        for candidate in ["python3", "python", "py -3"]:
            self.assertIn(candidate, runner)


class PortableHookRunnerTests(unittest.TestCase):
    def run_guard(self, command):
        if os.name == "nt":
            runner_command = r".cursor\hooks\run_hook.cmd .cursor\hooks\guard_shell.py"
        else:
            runner_command = "/bin/sh .cursor/hooks/run_hook.cmd .cursor/hooks/guard_shell.py"
        return subprocess.run(
            runner_command,
            input=json.dumps({"command": command}),
            text=True,
            capture_output=True,
            cwd=ROOT,
            shell=True,
            timeout=10,
        )

    def test_runner_executes_safe_hook(self):
        result = self.run_guard("whoami")
        self.assertEqual(result.returncode, 0, result.stderr)
        self.assertEqual(json.loads(result.stdout)["permission"], "allow")

    def test_runner_preserves_hook_denial_exit_code(self):
        result = self.run_guard("git reset --hard HEAD")
        self.assertEqual(result.returncode, 2, result.stderr)
        self.assertEqual(json.loads(result.stdout)["permission"], "deny")

    @unittest.skipIf(os.name == "nt", "POSIX fallback branch only")
    def test_posix_runner_falls_back_to_python_when_python3_is_absent(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            interpreter = Path(temp_dir) / "python"
            interpreter.symlink_to(sys.executable)
            env = os.environ.copy()
            env["PATH"] = temp_dir
            result = subprocess.run(
                "/bin/sh .cursor/hooks/run_hook.cmd .cursor/hooks/guard_shell.py",
                input=json.dumps({"command": "whoami"}),
                text=True,
                capture_output=True,
                cwd=ROOT,
                env=env,
                shell=True,
                timeout=10,
            )
            self.assertEqual(result.returncode, 0, result.stderr)
            self.assertEqual(json.loads(result.stdout)["permission"], "allow")


class PreflightTests(unittest.TestCase):
    def test_preflight_reports_machine_readable_capabilities(self):
        result = subprocess.run(
            [sys.executable, "scripts/ai-team/preflight.py", "--json"],
            text=True,
            capture_output=True,
            cwd=ROOT,
            timeout=15,
        )
        self.assertIn(result.returncode, {0, 1}, result.stderr)
        report = json.loads(result.stdout)
        self.assertEqual(report["hooks_config"]["status"], "pass")
        self.assertEqual(report["guard_hook"]["status"], "pass")
        self.assertEqual(report["project_cli"]["status"], "pass")
        self.assertEqual(report["global_allowlist"]["status"], "manual")
        self.assertEqual(report["execution_surface"]["status"], "manual")
        if report["platform"] == "windows-native":
            self.assertEqual(report["readonly_sandbox"]["status"], "skip")
        else:
            self.assertEqual(report["readonly_sandbox"]["status"], "expected")


class AuditEventTests(unittest.TestCase):
    def run_audit(self, stdin_data, env=None):
        merged = os.environ.copy()
        if env:
            merged.update(env)
        if isinstance(stdin_data, str):
            stdin_data = stdin_data.encode("utf-8")
        return subprocess.run(
            [sys.executable, str(CURSOR / "hooks" / "audit_event.py")],
            input=stdin_data,
            capture_output=True,
            cwd=ROOT,
            env=merged,
            timeout=10,
        )

    def test_valid_json_payload_is_logged_intact(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            result = self.run_audit(
                json.dumps(
                    {
                        "hook_event_name": "beforeShellExecution",
                        "command": "whoami",
                    }
                ),
                env={"CURSOR_PROJECT_DIR": temp_dir},
            )
            self.assertEqual(result.returncode, 0, result.stderr)
            self.assertEqual(json.loads(result.stdout.decode("utf-8"))["permission"], "allow")
            log_path = Path(temp_dir) / ".ai-team" / "logs" / "cursor-events.jsonl"
            record = json.loads(log_path.read_text(encoding="utf-8").strip())
            self.assertEqual(record["event"]["hook_event_name"], "beforeShellExecution")
            self.assertEqual(record["event"]["command"], "whoami")

    def test_bom_prefixed_json_is_parsed(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            payload = ("\ufeff" + json.dumps({"hook_event_name": "subagentStart"})).encode(
                "utf-8"
            )
            result = self.run_audit(payload, env={"CURSOR_PROJECT_DIR": temp_dir})
            self.assertEqual(result.returncode, 0, result.stderr)
            log_path = Path(temp_dir) / ".ai-team" / "logs" / "cursor-events.jsonl"
            record = json.loads(log_path.read_text(encoding="utf-8").strip())
            self.assertEqual(record["event"]["hook_event_name"], "subagentStart")

    def test_invalid_json_preserves_raw_instead_of_empty_string(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            result = self.run_audit(
                "not-json-but-should-be-kept",
                env={"CURSOR_PROJECT_DIR": temp_dir},
            )
            self.assertEqual(result.returncode, 0, result.stderr)
            log_path = Path(temp_dir) / ".ai-team" / "logs" / "cursor-events.jsonl"
            record = json.loads(log_path.read_text(encoding="utf-8").strip())
            self.assertEqual(record["event"]["raw"], "not-json-but-should-be-kept")


class GuardShellTests(unittest.TestCase):
    def run_guard(self, command, project_dir=None):
        env = os.environ.copy()
        if project_dir:
            env["CURSOR_PROJECT_DIR"] = str(project_dir)
        return subprocess.run(
            [sys.executable, str(CURSOR / "hooks" / "guard_shell.py")],
            input=json.dumps({"command": command}),
            text=True,
            capture_output=True,
            cwd=ROOT,
            env=env,
            timeout=10,
        )

    def initialize_project(self, target):
        (target / ".ai-team").mkdir(parents=True)
        (target / ".ai-team" / "project-profile.yaml").write_text(
            "release:\n  protected_branch: main\n", encoding="utf-8"
        )
        (target / "tracked.txt").write_text("initial\n", encoding="utf-8")
        commands = [
            ["git", "init", "-q"],
            ["git", "config", "user.name", "Guard Test"],
            ["git", "config", "user.email", "guard@example.invalid"],
            ["git", "add", "."],
            ["git", "commit", "-qm", "initial"],
            ["git", "branch", "-M", "main"],
        ]
        for command in commands:
            result = subprocess.run(
                command, cwd=target, text=True, capture_output=True, timeout=10
            )
            self.assertEqual(result.returncode, 0, result.stderr + result.stdout)

    def test_safe_command_is_allowed(self):
        result = self.run_guard("git status --short")
        self.assertEqual(result.returncode, 0, result.stderr)
        self.assertEqual(json.loads(result.stdout)["permission"], "allow")

    def test_hazardous_commands_are_denied(self):
        for command in [
            "git reset --hard HEAD~1",
            "git push origin main",
            "terraform destroy",
            "kubectl delete deployment app",
            "rm -rf /",
        ]:
            with self.subTest(command=command):
                result = self.run_guard(command)
                self.assertEqual(result.returncode, 2, result.stderr)
                self.assertEqual(json.loads(result.stdout)["permission"], "deny")

    def test_working_branch_commits_are_allowed_but_protected_branch_is_blocked(self):
        with tempfile.TemporaryDirectory(dir=ROOT) as temp_dir:
            target = Path(temp_dir)
            self.initialize_project(target)

            protected = self.run_guard("git add tracked.txt", target)
            self.assertEqual(protected.returncode, 2, protected.stderr)
            self.assertIn("protected branch", json.loads(protected.stdout)["message"])

            switch = subprocess.run(
                ["git", "switch", "-c", "work/WU-TEST"],
                cwd=target,
                text=True,
                capture_output=True,
                timeout=10,
            )
            self.assertEqual(switch.returncode, 0, switch.stderr + switch.stdout)
            working = self.run_guard(
                'git commit -m "feat(WU-TEST): coherent change"', target
            )
            self.assertEqual(working.returncode, 0, working.stderr)

            missing_work_unit = self.run_guard(
                'git commit -m "feat: missing governed scope"', target
            )
            self.assertEqual(missing_work_unit.returncode, 2, missing_work_unit.stderr)
            self.assertIn("type(WU-ID)", json.loads(missing_work_unit.stdout)["message"])

    def test_framework_source_allows_conventional_commits(self):
        with tempfile.TemporaryDirectory(dir=ROOT) as temp_dir:
            target = Path(temp_dir)
            (target / ".ai-team").mkdir(parents=True)
            (target / ".ai-team" / "project-profile.yaml").write_text(
                "project:\n  repository_kind: framework_source\n"
                "release:\n  protected_branch: main\n",
                encoding="utf-8",
            )
            (target / "tracked.txt").write_text("initial\n", encoding="utf-8")
            commands = [
                ["git", "init", "-q"],
                ["git", "config", "user.name", "Guard Test"],
                ["git", "config", "user.email", "guard@example.invalid"],
                ["git", "add", "."],
                ["git", "commit", "-qm", "initial"],
                ["git", "branch", "-M", "main"],
                ["git", "switch", "-c", "renov/test"],
            ]
            for command in commands:
                result = subprocess.run(
                    command, cwd=target, text=True, capture_output=True, timeout=10
                )
                self.assertEqual(result.returncode, 0, result.stderr + result.stdout)

            working = self.run_guard('git commit -m "feat: fabrication change"', target)
            self.assertEqual(working.returncode, 0, working.stderr)

    def test_history_rewrite_is_not_an_autonomous_path(self):
        for command in ["git commit --amend --no-edit", "git rebase main"]:
            with self.subTest(command=command):
                result = self.run_guard(command)
                self.assertEqual(result.returncode, 2, result.stderr)
                self.assertIn("History rewriting", json.loads(result.stdout)["message"])


class InstallerCliIntegrationTests(unittest.TestCase):
    def run_command(self, args, cwd=ROOT):
        return subprocess.run(
            args,
            cwd=cwd,
            text=True,
            capture_output=True,
            timeout=30,
        )

    def initialize_git(self, target):
        commands = [
            ["git", "init", "-q"],
            ["git", "config", "user.name", "Framework Test"],
            ["git", "config", "user.email", "framework-test@example.invalid"],
            ["git", "add", "."],
            ["git", "commit", "-qm", "test fixture"],
        ]
        for command in commands:
            result = self.run_command(command, cwd=target)
            self.assertEqual(result.returncode, 0, result.stderr + result.stdout)

    def install_target(self, target, project_id="cli-test", project_name="CLI Test"):
        install = self.run_command(
            [
                sys.executable,
                "tools/install.py",
                "--target",
                str(target),
                "--project-id",
                project_id,
                "--project-name",
                project_name,
            ]
        )
        self.assertEqual(install.returncode, 0, install.stderr + install.stdout)
        return install

    def test_install_update_and_validation_keep_both_cursor_modes(self):
        with tempfile.TemporaryDirectory(dir=ROOT) as temp_dir:
            target = Path(temp_dir) / "target-project"
            self.install_target(target)
            self.assertTrue((target / ".cursor" / "permissions.json").is_file())
            self.assertTrue((target / ".cursor" / "cli.json").is_file())
            installed_runner = target / ".cursor" / "hooks" / "run_hook.cmd"
            self.assertTrue(installed_runner.is_file())
            if os.name != "nt":
                from adapters.cursor.runtime.checks import probe_hook

                guard_ok, guard_detail = probe_hook(
                    target, "guard_shell.py", {"command": "whoami"}
                )
                self.assertTrue(guard_ok, guard_detail)
            self.assertTrue((target / "scripts" / "ai-team" / "preflight.py").is_file())

            validate = self.run_command(
                [sys.executable, "scripts/ai-team/validate.py"], cwd=target
            )
            self.assertEqual(validate.returncode, 0, validate.stderr + validate.stdout)

            profile_path = target / ".ai-team" / "project-profile.yaml"
            profile_before_update = profile_path.read_text(encoding="utf-8")
            cli_path = target / ".cursor" / "cli.json"
            cli_config = json.loads(cli_path.read_text(encoding="utf-8"))
            cli_config["permissions"]["allow"].append("Shell(test-local-override)")
            cli_path.write_text(json.dumps(cli_config), encoding="utf-8")
            self.initialize_git(target)

            # .cursor/ is fully generated/compiler-owned output (never hand-edited
            # per its own documentation); a local edit is therefore local drift,
            # which --update now refuses by default unless --force is passed.
            update = self.run_command(
                [
                    sys.executable,
                    "tools/install.py",
                    "--target",
                    str(target),
                    "--update",
                    "--force",
                ]
            )
            self.assertEqual(update.returncode, 0, update.stderr + update.stdout)
            self.assertEqual(profile_path.read_text(encoding="utf-8"), profile_before_update)
            updated_cli = json.loads(cli_path.read_text(encoding="utf-8"))
            self.assertNotIn(
                "Shell(test-local-override)", updated_cli["permissions"]["allow"]
            )
            self.assertTrue((target / ".cursor" / "permissions.json").is_file())
            manifest = json.loads(
                (target / ".ai-team" / "framework-version.json").read_text(
                    encoding="utf-8"
                )
            )
            self.assertEqual(manifest["version"], "0.7.0")
            self.assertIn(".cursor/hooks.json", manifest["managed_files"])

    def test_propose_allowlist_derives_tokens_from_declared_commands_only(self):
        with tempfile.TemporaryDirectory(dir=ROOT) as temp_dir:
            target = Path(temp_dir) / "target-project"
            self.install_target(target, "allowlist-test", "Allowlist Test")
            self.assertTrue((target / "scripts" / "ai-team" / "propose_allowlist.py").is_file())

            profile_path = target / ".ai-team" / "project-profile.yaml"
            profile = yaml.safe_load(profile_path.read_text(encoding="utf-8"))
            profile["commands"] = {
                "setup": "make install",
                "build": "make check",
                "lint": "make lint",
                "typecheck": "make typecheck",
                "unit_test": "make test",
                "integration_test": "make test",
                "e2e_test": "cd frontend && pnpm test:e2e",
            }
            profile_path.write_text(yaml.safe_dump(profile, sort_keys=False), encoding="utf-8")

            result = self.run_command(
                [sys.executable, "scripts/ai-team/propose_allowlist.py", "--json"],
                cwd=target,
            )
            self.assertEqual(result.returncode, 0, result.stderr + result.stdout)
            report = json.loads(result.stdout)

            self.assertIn("make install", report["terminal_allowlist_additions"])
            self.assertIn("make check", report["terminal_allowlist_additions"])
            self.assertIn("make lint", report["terminal_allowlist_additions"])
            self.assertIn("make typecheck", report["terminal_allowlist_additions"])
            self.assertIn("make test", report["terminal_allowlist_additions"])
            # integration_test duplicates unit_test's "make test" - proposed once, not twice.
            self.assertEqual(report["terminal_allowlist_additions"].count("make test"), 1)

            self.assertIn("Shell(make:install*)", report["cli_allow_additions"])
            self.assertIn("Shell(make:check*)", report["cli_allow_additions"])
            self.assertIn("Shell(make:test*)", report["cli_allow_additions"])

            # No tool-level wildcard is ever emitted.
            joined = " ".join(report["cli_allow_additions"])
            self.assertNotIn("Shell(make:*)", joined)

            # A compound command (shell operators) has no single binary to key a
            # Shell(tool:args) token on: it must be flagged for manual review, not
            # guessed at (e.g. never "Shell(cd:frontend && pnpm test:e2e*)").
            self.assertIn("cd frontend && pnpm test:e2e", report["terminal_allowlist_additions"])
            self.assertFalse(
                any("cd" in token for token in report["cli_allow_additions"]),
                report["cli_allow_additions"],
            )
            manual_review_commands = {
                item["command"] for item in report["cli_allow_needs_manual_review"]
            }
            self.assertIn("cd frontend && pnpm test:e2e", manual_review_commands)

            cli_path = target / ".cursor" / "cli.json"
            cli_config = json.loads(cli_path.read_text(encoding="utf-8"))
            cli_config["permissions"]["allow"].append("Shell(make:check*)")
            cli_path.write_text(json.dumps(cli_config), encoding="utf-8")

            rerun = self.run_command(
                [sys.executable, "scripts/ai-team/propose_allowlist.py", "--json"],
                cwd=target,
            )
            self.assertEqual(rerun.returncode, 0, rerun.stderr + rerun.stdout)
            rerun_report = json.loads(rerun.stdout)
            # Already-present entries are not re-proposed.
            self.assertNotIn("Shell(make:check*)", rerun_report["cli_allow_additions"])
            self.assertIn("Shell(make:install*)", rerun_report["cli_allow_additions"])

    def test_scripts_follow_communication_language_for_their_own_output(self):
        with tempfile.TemporaryDirectory(dir=ROOT) as temp_dir:
            target = Path(temp_dir) / "target-project"
            self.install_target(target, "lang-test", "Lang Test")
            profile_path = target / ".ai-team" / "project-profile.yaml"
            base_text = profile_path.read_text(encoding="utf-8")

            english_status = self.run_command(
                [sys.executable, "scripts/ai-team/status.py"], cwd=target
            )
            self.assertEqual(english_status.returncode, 0, english_status.stderr)
            self.assertIn("Projet :", english_status.stdout)

            profile_path.write_text(
                base_text.replace("language: français", "language: english"),
                encoding="utf-8",
            )
            english_status = self.run_command(
                [sys.executable, "scripts/ai-team/status.py"], cwd=target
            )
            self.assertEqual(english_status.returncode, 0, english_status.stderr)
            self.assertIn("Project:", english_status.stdout)
            self.assertNotIn("Projet :", english_status.stdout)

            profile_path.write_text(
                base_text,
                encoding="utf-8",
            )
            french_status = self.run_command(
                [sys.executable, "scripts/ai-team/status.py"], cwd=target
            )
            self.assertEqual(french_status.returncode, 0, french_status.stderr)
            self.assertIn("Projet :", french_status.stdout)
            self.assertNotIn("Project:", french_status.stdout)
            # Enum values and YAML-facing content are never translated.
            self.assertIn("not_required", french_status.stdout)

            missing_wu = self.run_command(
                [sys.executable, "scripts/ai-team/check_done.py", "WU-DOES-NOT-EXIST"],
                cwd=target,
            )
            self.assertEqual(missing_wu.returncode, 2)
            self.assertIn("introuvable", missing_wu.stdout)

            validate = self.run_command(
                [sys.executable, "scripts/ai-team/validate.py"], cwd=target
            )
            self.assertIn("Validation Governed AI Team", validate.stdout)
            self.assertIn("erreur(s)", validate.stdout)
            # Error/warning bodies stay in English regardless of project language.
            self.assertIn("WARN  Project command", validate.stdout)

    def test_status_surfaces_open_human_checkpoints_deduped_by_work_unit(self):
        with tempfile.TemporaryDirectory(dir=ROOT) as temp_dir:
            target = Path(temp_dir) / "target-project"
            self.install_target(target, "checkpoint-test", "Checkpoint Test")
            events_dir = target / ".ai-team" / "events"
            events_dir.mkdir(parents=True, exist_ok=True)

            def write_event(name, work_unit, status, why):
                (events_dir / f"{name}.yaml").write_text(
                    yaml.safe_dump(
                        {
                            "id": name,
                            "type": "STATUS",
                            "work_unit": work_unit,
                            "created_at": "2026-08-26T12:00:00+00:00",
                            "created_by_role": "qa-test",
                            "summary": "smoke",
                            "details": {
                                "human_checkpoint": {
                                    "command": "npm run dev",
                                    "why": why,
                                    "states_to_check": ["initial_state"],
                                }
                            },
                            "affected_nodes": [work_unit],
                            "requires_human": False,
                            "status": status,
                        },
                        sort_keys=False,
                    ),
                    encoding="utf-8",
                )

            write_event("EVT-OPEN-FE1-OLD", "WU-FE-0001", "open", "first pass")
            write_event("EVT-OPEN-FE1-NEW", "WU-FE-0001", "open", "second pass")
            write_event("EVT-CLOSED-FE2", "WU-FE-0002", "closed", "already reviewed")
            write_event("EVT-NO-CHECKPOINT", "WU-BE-0001", "open", "n/a")
            (events_dir / "EVT-NO-CHECKPOINT.yaml").write_text(
                yaml.safe_dump(
                    {
                        "id": "EVT-NO-CHECKPOINT",
                        "type": "STATUS",
                        "work_unit": "WU-BE-0001",
                        "created_at": "2026-08-26T12:00:00+00:00",
                        "created_by_role": "backend-developer",
                        "summary": "no UI involved",
                        "details": {},
                        "affected_nodes": [],
                        "requires_human": False,
                        "status": "open",
                    },
                    sort_keys=False,
                ),
                encoding="utf-8",
            )

            result = self.run_command(
                [sys.executable, "scripts/ai-team/status.py"], cwd=target
            )
            self.assertEqual(result.returncode, 0, result.stderr)
            self.assertIn("Points de controle visuels disponibles : 1", result.stdout)
            self.assertIn("WU-FE-0001", result.stdout)
            self.assertNotIn("WU-FE-0002", result.stdout)
            self.assertNotIn("WU-BE-0001", result.stdout)

    def test_read_docx_extracts_paragraphs_and_table_rows(self):
        try:
            import docx
        except ModuleNotFoundError:
            self.skipTest("python-docx not installed")
        with tempfile.TemporaryDirectory(dir=ROOT) as temp_dir:
            target = Path(temp_dir) / "target-project"
            self.install_target(target, "docx-test", "Docx Test")

            document = docx.Document()
            document.add_paragraph("La VersionExercice devient immuable des l'attribution.")
            table = document.add_table(rows=1, cols=3)
            table.rows[0].cells[0].text = "INV-004"
            table.rows[0].cells[1].text = "INV"
            table.rows[0].cells[2].text = "Une version deja utilisee est immuable."
            docx_path = target / "sample.docx"
            document.save(str(docx_path))

            full = self.run_command(
                [sys.executable, "scripts/ai-team/read_docx.py", "sample.docx"],
                cwd=target,
            )
            self.assertEqual(full.returncode, 0, full.stderr)
            self.assertIn("VersionExercice devient immuable", full.stdout)
            self.assertIn("INV-004 | INV | Une version deja utilisee est immuable.", full.stdout)

            grepped = self.run_command(
                [
                    sys.executable,
                    "scripts/ai-team/read_docx.py",
                    "sample.docx",
                    "--grep",
                    "INV-004",
                    "--context",
                    "10",
                ],
                cwd=target,
            )
            self.assertEqual(grepped.returncode, 0, grepped.stderr)
            self.assertIn("INV-004", grepped.stdout)
            self.assertNotIn("VersionExercice", grepped.stdout)

            missing = self.run_command(
                [sys.executable, "scripts/ai-team/read_docx.py", "does-not-exist.docx"],
                cwd=target,
            )
            self.assertEqual(missing.returncode, 2)

    def test_validation_rejects_global_only_settings_in_project_cli_config(self):
        with tempfile.TemporaryDirectory(dir=ROOT) as temp_dir:
            target = Path(temp_dir) / "target-project"
            self.install_target(target, "global-setting-test", "Global Setting Test")
            cli_path = target / ".cursor" / "cli.json"
            cli_config = json.loads(cli_path.read_text(encoding="utf-8"))
            cli_config["approvalMode"] = "allowlist"
            cli_path.write_text(json.dumps(cli_config), encoding="utf-8")

            validate = self.run_command(
                [sys.executable, "scripts/ai-team/validate.py"], cwd=target
            )
            self.assertEqual(validate.returncode, 1)
            self.assertIn("only permissions can be configured at project level", validate.stdout)

    def test_validation_rejects_malformed_cli_configuration(self):
        with tempfile.TemporaryDirectory(dir=ROOT) as temp_dir:
            target = Path(temp_dir) / "target-project"
            self.install_target(target, "invalid-cli-test", "Invalid CLI Test")
            (target / ".cursor" / "cli.json").write_text("{not-json", encoding="utf-8")

            validate = self.run_command(
                [sys.executable, "scripts/ai-team/validate.py"], cwd=target
            )
            self.assertEqual(validate.returncode, 1)
            self.assertIn("Invalid JSON .cursor", validate.stdout)

    def test_update_dry_run_is_read_only_even_without_site_packages(self):
        with tempfile.TemporaryDirectory(dir=ROOT) as temp_dir:
            target = Path(temp_dir) / "target-project"
            self.install_target(target)
            cli_path = target / ".cursor" / "cli.json"
            before = cli_path.read_bytes()
            result = self.run_command(
                [
                    sys.executable,
                    "-S",
                    "tools/install.py",
                    "--target",
                    str(target),
                    "--update",
                    "--dry-run",
                ]
            )
            self.assertEqual(result.returncode, 0, result.stderr + result.stdout)
            self.assertIn("DRY-RUN: no file was modified", result.stdout)
            self.assertEqual(cli_path.read_bytes(), before)

    @unittest.skipIf(os.name == "nt", "WSL/Linux target-venv selection test")
    def test_update_started_without_site_packages_uses_target_venv_for_validation(self):
        with tempfile.TemporaryDirectory(dir=ROOT) as temp_dir:
            target = Path(temp_dir) / "target-project"
            self.install_target(target)
            marker = target / ".ai-team" / "framework-version.json"
            marker.unlink()
            target_python = target / ".venv" / "bin" / "python"
            target_python.parent.mkdir(parents=True)
            target_python.write_text(
                f'#!/bin/sh\nexec "{sys.executable}" "$@"\n', encoding="utf-8"
            )
            target_python.chmod(0o755)
            self.initialize_git(target)

            result = self.run_command(
                [
                    sys.executable,
                    "-S",
                    "tools/install.py",
                    "--target",
                    str(target),
                    "--update",
                ]
            )
            self.assertEqual(result.returncode, 0, result.stderr + result.stdout)
            self.assertIn("Post-update validation: PASS", result.stdout)
            self.assertTrue(marker.is_file())

    def test_update_aborts_on_dirty_target_before_writing(self):
        with tempfile.TemporaryDirectory(dir=ROOT) as temp_dir:
            target = Path(temp_dir) / "target-project"
            self.install_target(target)
            self.initialize_git(target)
            # Dirty the worktree via a non-managed, project-owned file so this
            # test exercises the git-dirty guard in isolation from the
            # separate local-drift guard (which fires first, with its own
            # message, when the edited file is itself framework-managed).
            marker_path = target / "app_notes.txt"
            marker_path.write_text("local notes\n", encoding="utf-8")
            cli_path = target / ".cursor" / "cli.json"
            original = cli_path.read_text(encoding="utf-8")
            result = self.run_command(
                [sys.executable, "tools/install.py", "--target", str(target), "--update"]
            )
            self.assertEqual(result.returncode, 2, result.stderr + result.stdout)
            self.assertIn("target must be a clean", result.stdout)
            self.assertEqual(cli_path.read_text(encoding="utf-8"), original)

    def test_update_rejects_unknown_future_version_before_writing(self):
        with tempfile.TemporaryDirectory(dir=ROOT) as temp_dir:
            target = Path(temp_dir) / "target-project"
            self.install_target(target)
            marker = target / ".ai-team" / "framework-version.json"
            record_path = target / ".ai-team" / "installation-record.json"
            payload = json.loads(marker.read_text(encoding="utf-8"))
            payload["version"] = "99.0.0"
            marker.write_text(json.dumps(payload), encoding="utf-8")
            if record_path.is_file():
                record = json.loads(record_path.read_text(encoding="utf-8"))
                record["core"]["version"] = "99.0.0"
                record["distribution"]["version"] = "99.0.0"
                record["adapters"][0]["version"] = "99.0.0"
                record_path.write_text(json.dumps(record), encoding="utf-8")
            cli_path = target / ".cursor" / "cli.json"
            cli_before = cli_path.read_bytes()
            self.initialize_git(target)

            result = self.run_command(
                [sys.executable, "tools/install.py", "--target", str(target), "--update"]
            )
            self.assertEqual(result.returncode, 2, result.stderr + result.stdout)
            self.assertIn("No safe migration path", result.stdout)
            self.assertEqual(cli_path.read_bytes(), cli_before)
            self.assertEqual(
                json.loads(marker.read_text(encoding="utf-8"))["version"], "99.0.0"
            )
            if record_path.is_file():
                record_after = json.loads(record_path.read_text(encoding="utf-8"))
                self.assertEqual(record_after["core"]["version"], "99.0.0")

    def test_update_activates_new_constitution_only_between_cycles(self):
        with tempfile.TemporaryDirectory(dir=ROOT) as temp_dir:
            target = Path(temp_dir) / "target-project"
            self.install_target(target)
            marker = target / ".ai-team" / "framework-version.json"
            marker_payload = json.loads(marker.read_text(encoding="utf-8"))
            marker_payload["version"] = "0.3.0"
            marker.write_text(json.dumps(marker_payload), encoding="utf-8")
            state_path = target / ".ai-team" / "state" / "project-state.yaml"
            state = yaml.safe_load(state_path.read_text(encoding="utf-8"))
            state["constitution_version"] = "1.0.0"
            state["phase"] = "not_compiled"
            state_path.write_text(
                yaml.safe_dump(state, sort_keys=False), encoding="utf-8"
            )
            self.initialize_git(target)

            result = self.run_command(
                [sys.executable, "tools/install.py", "--target", str(target), "--update"]
            )
            self.assertEqual(result.returncode, 0, result.stderr + result.stdout)
            self.assertIn("constitution_version 1.0.0 -> 1.2.0", result.stdout)
            migrated = yaml.safe_load(state_path.read_text(encoding="utf-8"))
            self.assertEqual(migrated["constitution_version"], "1.2.0")

    def test_update_refuses_constitution_change_during_active_cycle(self):
        with tempfile.TemporaryDirectory(dir=ROOT) as temp_dir:
            target = Path(temp_dir) / "target-project"
            self.install_target(target)
            marker = target / ".ai-team" / "framework-version.json"
            marker_payload = json.loads(marker.read_text(encoding="utf-8"))
            marker_payload["version"] = "0.3.0"
            marker.write_text(json.dumps(marker_payload), encoding="utf-8")
            state_path = target / ".ai-team" / "state" / "project-state.yaml"
            state = yaml.safe_load(state_path.read_text(encoding="utf-8"))
            state["constitution_version"] = "1.0.0"
            state["phase"] = "execution"
            state_path.write_text(
                yaml.safe_dump(state, sort_keys=False), encoding="utf-8"
            )
            self.initialize_git(target)
            before = state_path.read_bytes()

            result = self.run_command(
                [sys.executable, "tools/install.py", "--target", str(target), "--update"]
            )
            self.assertEqual(result.returncode, 2, result.stderr + result.stdout)
            self.assertIn("Constitution 1.0.0 is frozen", result.stdout)
            self.assertEqual(state_path.read_bytes(), before)
            self.assertEqual(
                json.loads(marker.read_text(encoding="utf-8"))["version"], "0.3.0"
            )

    def test_update_force_constitution_update_bypasses_freeze_and_logs_event(self):
        with tempfile.TemporaryDirectory(dir=ROOT) as temp_dir:
            target = Path(temp_dir) / "target-project"
            self.install_target(target)
            marker = target / ".ai-team" / "framework-version.json"
            marker_payload = json.loads(marker.read_text(encoding="utf-8"))
            marker_payload["version"] = "0.3.0"
            marker.write_text(json.dumps(marker_payload), encoding="utf-8")
            state_path = target / ".ai-team" / "state" / "project-state.yaml"
            state = yaml.safe_load(state_path.read_text(encoding="utf-8"))
            state["constitution_version"] = "1.0.0"
            state["phase"] = "execution"
            state_path.write_text(
                yaml.safe_dump(state, sort_keys=False), encoding="utf-8"
            )
            self.initialize_git(target)
            events_dir = target / ".ai-team" / "events"
            before_events = set(events_dir.glob("*.yaml")) if events_dir.exists() else set()

            result = self.run_command(
                [
                    sys.executable,
                    "tools/install.py",
                    "--target",
                    str(target),
                    "--update",
                    "--force-constitution-update",
                ]
            )
            self.assertEqual(result.returncode, 0, result.stderr + result.stdout)
            self.assertIn("WARNING: forcing Constitution", result.stdout)
            migrated = yaml.safe_load(state_path.read_text(encoding="utf-8"))
            self.assertEqual(migrated["constitution_version"], "1.2.0")

            after_events = set(events_dir.glob("*.yaml"))
            new_events = after_events - before_events
            self.assertEqual(len(new_events), 1)
            event = yaml.safe_load(new_events.pop().read_text(encoding="utf-8"))
            self.assertEqual(event["type"], "CONTRACT_CHANGE")
            self.assertTrue(event["requires_human"])
            self.assertEqual(event["status"], "open")
            self.assertEqual(event["details"]["old_constitution_version"], "1.0.0")
            self.assertEqual(event["details"]["new_constitution_version"], "1.2.0")
            self.assertEqual(event["details"]["phase_at_override"], "execution")

    def test_force_constitution_update_requires_update_flag(self):
        with tempfile.TemporaryDirectory(dir=ROOT) as temp_dir:
            target = Path(temp_dir) / "target-project"
            result = self.run_command(
                [
                    sys.executable,
                    "tools/install.py",
                    "--target",
                    str(target),
                    "--force-constitution-update",
                    "--project-id",
                    "demo",
                    "--project-name",
                    "Demo",
                ]
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("require --update", result.stderr + result.stdout)

    def test_legacy_update_migrates_acceptance_and_keeps_backup(self):
        with tempfile.TemporaryDirectory(dir=ROOT) as temp_dir:
            target = Path(temp_dir) / "target-project"
            self.install_target(target)
            (target / ".ai-team" / "framework-version.json").unlink()
            record_path = target / ".ai-team" / "installation-record.json"
            if record_path.is_file():
                record_path.unlink()
            acceptance = target / ".ai-team" / "acceptance" / "ACC-LEGACY.yaml"
            acceptance.write_text(
                "id: ACC-LEGACY\n"
                "revision: 1\n"
                "created_at: '2026-01-01T00:00:00+00:00'\n"
                "updated_at: '2026-01-01T00:00:00+00:00'\n"
                "scenarios: []\n"
                "human_result:\n"
                "  status: accepted\n",
                encoding="utf-8",
            )
            cli_path = target / ".cursor" / "cli.json"
            cli_config = json.loads(cli_path.read_text(encoding="utf-8"))
            cli_config["permissions"]["allow"].append("Shell(legacy-override)")
            cli_path.write_text(json.dumps(cli_config), encoding="utf-8")
            self.initialize_git(target)

            result = self.run_command(
                [sys.executable, "tools/install.py", "--target", str(target), "--update"]
            )
            self.assertEqual(result.returncode, 0, result.stderr + result.stdout)
            self.assertIn("Post-update validation: PASS", result.stdout)
            self.assertIn("status: passed", acceptance.read_text(encoding="utf-8"))
            backup = (
                target
                / ".ai-team"
                / "migration-backups"
                / "acceptance-status-passed"
                / ".ai-team"
                / "acceptance"
                / "ACC-LEGACY.yaml"
            )
            self.assertIn("status: accepted", backup.read_text(encoding="utf-8"))
            updated_cli = json.loads(cli_path.read_text(encoding="utf-8"))
            self.assertNotIn("Shell(legacy-override)", updated_cli["permissions"]["allow"])
            manifest = json.loads(
                (target / ".ai-team" / "framework-version.json").read_text(
                    encoding="utf-8"
                )
            )
            self.assertEqual(manifest["version"], "0.7.0")

    def test_failed_post_update_validation_rolls_back_all_touched_files(self):
        with tempfile.TemporaryDirectory(dir=ROOT) as temp_dir:
            target = Path(temp_dir) / "target-project"
            self.install_target(target)
            marker = target / ".ai-team" / "framework-version.json"
            marker.unlink()
            record_path = target / ".ai-team" / "installation-record.json"
            if record_path.is_file():
                record_path.unlink()
            acceptance = target / ".ai-team" / "acceptance" / "ACC-INVALID.yaml"
            acceptance.write_text(
                "id: ACC-INVALID\n"
                "revision: 1\n"
                "created_at: '2026-01-01T00:00:00+00:00'\n"
                "updated_at: '2026-01-01T00:00:00+00:00'\n"
                "scenarios: []\n"
                "human_result:\n"
                "  status: invalid\n",
                encoding="utf-8",
            )
            cli_path = target / ".cursor" / "cli.json"
            cli_config = json.loads(cli_path.read_text(encoding="utf-8"))
            cli_config["permissions"]["allow"].append("Shell(must-survive-rollback)")
            cli_path.write_text(json.dumps(cli_config), encoding="utf-8")
            self.initialize_git(target)

            result = self.run_command(
                [sys.executable, "tools/install.py", "--target", str(target), "--update"]
            )
            self.assertEqual(result.returncode, 1, result.stderr + result.stdout)
            self.assertIn("Update rolled back", result.stdout)
            self.assertFalse(marker.exists())
            rolled_back_cli = json.loads(cli_path.read_text(encoding="utf-8"))
            self.assertIn(
                "Shell(must-survive-rollback)", rolled_back_cli["permissions"]["allow"]
            )


class MigrationTests(unittest.TestCase):
    def run_migration(self, target, apply=False):
        command = [
            sys.executable,
            "scripts/ai-team/migrate.py",
            "--target",
            str(target),
        ]
        if apply:
            command.append("--apply")
        return subprocess.run(
            command,
            cwd=ROOT,
            text=True,
            capture_output=True,
            timeout=15,
        )

    def test_standalone_migration_is_dry_run_then_idempotent_apply(self):
        with tempfile.TemporaryDirectory(dir=ROOT) as temp_dir:
            target = Path(temp_dir)
            acceptance_dir = target / ".ai-team" / "acceptance"
            acceptance_dir.mkdir(parents=True)
            acceptance = acceptance_dir / "ACC-WU-BE-0004.yaml"
            original = (
                "id: ACC-WU-BE-0004\n"
                "scenarios: []\n"
                "human_result:\n"
                "  status: 'accepted' # legacy value\n"
            )
            acceptance.write_text(original, encoding="utf-8")

            dry_run = self.run_migration(target)
            self.assertEqual(dry_run.returncode, 0, dry_run.stderr + dry_run.stdout)
            self.assertIn("MIGRATE", dry_run.stdout)
            self.assertEqual(acceptance.read_text(encoding="utf-8"), original)

            applied = self.run_migration(target, apply=True)
            self.assertEqual(applied.returncode, 0, applied.stderr + applied.stdout)
            self.assertIn("status: 'passed' # legacy value", acceptance.read_text())
            second = self.run_migration(target, apply=True)
            self.assertEqual(second.returncode, 0, second.stderr + second.stdout)
            self.assertIn("No project-data migration required", second.stdout)

    def test_migration_only_changes_status_inside_human_result(self):
        with tempfile.TemporaryDirectory(dir=ROOT) as temp_dir:
            target = Path(temp_dir)
            acceptance_dir = target / ".ai-team" / "acceptance"
            acceptance_dir.mkdir(parents=True)
            acceptance = acceptance_dir / "ACC-SCOPED.yaml"
            acceptance.write_text(
                "status: accepted\n"
                "machine_result:\n"
                "  status: accepted\n"
                "human_result:\n"
                "  status: accepted\n",
                encoding="utf-8",
            )

            result = self.run_migration(target, apply=True)
            self.assertEqual(result.returncode, 0, result.stderr + result.stdout)
            self.assertEqual(
                acceptance.read_text(encoding="utf-8"),
                "status: accepted\n"
                "machine_result:\n"
                "  status: accepted\n"
                "human_result:\n"
                "  status: passed\n",
            )


if __name__ == "__main__":
    unittest.main()
